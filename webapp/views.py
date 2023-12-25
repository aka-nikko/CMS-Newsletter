from .models import Newsletter, Submission, Quarter, Topic
from .forms import NewsletterForm, CreateUserForm, SubmissionForm, SearchForm, AddTopicForm
from django.contrib import messages
from django.utils import timezone
from django.template.defaultfilters import striptags
from django.shortcuts import get_object_or_404, render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from docx import Document
from django.http import HttpResponse
import html2text

# Create your views here.
def indexView(request):
    try:
        data = Newsletter.objects.latest('upload_time')
    except Newsletter.DoesNotExist:
        data = None
    context = {'data':data}
    return render(request,'index.html', context)

def registerView(request):
    if request.user.is_authenticated:
        return redirect("home")
    if request.method == 'POST':
        form = CreateUserForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Account Created Successfully')
            return redirect("home")
    else:
        form = CreateUserForm()
    context = {'form':form}
    return render(request,'register.html',context)

def loginView(request):
    if request.user.is_authenticated:
        return redirect("home")
    else:
        if request.method == 'POST':
            username = request.POST.get('username')
            password = request.POST.get('password')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                remember_me_values = request.POST.getlist('remember-me')
                if remember_me_values:
                    request.session.set_expiry(2592000)
                else:
                    request.session.set_expiry(0)
                messages.success(request, 'Successfully Logged In !!')
                return redirect("home")
            else:
                messages.error(request, 'Wrong Email or Password !!')
                return redirect("login")
    return render(request,"login.html")

@login_required(login_url='home')
def addtopicView(request):
    form = AddTopicForm()
    quarter = Quarter.objects.all()
    topic = Topic.objects.all()

    if request.method=="POST":
        form = AddTopicForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'New Topic Added')
            return redirect("addtopic")

    context = {'form':form,'quarter':quarter,'topic':topic}
    return render(request,"addtopic.html", context)

@login_required(login_url='home')
def newsubmissionView(request):
    form = SubmissionForm()
    quarter = Quarter.objects.all()
    topic = Topic.objects.all()

    if request.method=="POST":
        form = SubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            save = form.save(commit=False)
            save.user = request.user
            save.save()
            messages.success(request, 'New Article Submitted')
            return redirect("submissions")

    context = {'form':form,'quarter':quarter,'topic':topic}
    return render(request,"newsubmission.html", context)

@login_required(login_url='home')
def submissionsView(request):
    quarters = Quarter.objects.all()
    selected_quarter_id = request.GET.get('quarter')

    if selected_quarter_id:
        if request.user.is_superuser:
            data = Submission.objects.filter(quarter_id=selected_quarter_id).order_by("upload_time")
        else:
            data = Submission.objects.filter(user=request.user, quarter_id=selected_quarter_id).order_by("upload_time")
    else:
        if request.user.is_superuser:
            data = Submission.objects.all().order_by("upload_time")
        else:
            data = Submission.objects.filter(user=request.user).order_by("upload_time")

    if request.method == "POST":
        if 'send_selected' in request.POST:
            selected_submission_ids = request.POST.getlist('selected_submissions')
            if selected_submission_ids:
                Submission.objects.filter(id__in=selected_submission_ids).update(
                    is_sent_for_editing=True,
                    sent_for_edit_time=timezone.now()
                )
                Submission.objects.exclude(id__in=selected_submission_ids).update(
                    is_sent_for_editing=False,
                    sent_for_edit_time=None
                )
                return redirect("editsubmissions")
            else:
                messages.error(request, 'No Submissions Selected For Editing')
                return redirect("submissions")
        if 'reset' in request.POST:
            Submission.objects.all().update(
                is_sent_for_editing=False,
                sent_for_edit_time=None
            )
            return redirect("submissions")

    context = {
        'data': data,
        'quarters': quarters,
        'selected_quarter_id': selected_quarter_id,
    }
    return render(request, 'submissions.html', context)

@login_required(login_url='home')
def submissionView(request, pk):
    data = Submission.objects.get(id=pk)

    if request.method == "POST":
        if 'send' in request.POST:
            submission_id = request.POST.get('submission_id')
            if submission_id:
                submission = get_object_or_404(Submission, id=submission_id)
                submission.is_sent_for_editing = True
                submission.save()
                return redirect("/submission/"+pk)
        if 'remove' in request.POST:
            submission_id = request.POST.get('submission_id')
            if submission_id:
                submission = get_object_or_404(Submission, id=submission_id)
                submission.is_sent_for_editing = False
                submission.save()
                return redirect("/submission/"+pk)
        if 'delete' in request.POST:
            data.delete()
            messages.error(request, 'Submission Deleted')
            return redirect("submissions")

    context = {'data':data}
    return render(request,'submission.html', context)

@login_required(login_url='home')
def submissionupdateView(request, pk):
    data = Submission.objects.get(id=pk)
    form = SubmissionForm(instance=data)
    quarter = Quarter.objects.all()
    topic = Topic.objects.all()
    is_editing = True

    if request.method == "POST":
        form = SubmissionForm(request.POST, request.FILES, instance=data)
        if form.is_valid():
            form.save()
            messages.success(request, 'Submission Updated')
            return redirect("/submission/"+pk)

    content = {'form': form, 'quarter': quarter, 'topic': topic, 'is_editing': is_editing}
    return render(request, "newsubmission.html", content)

@login_required(login_url='home')
def editsubmissionsView(request):
    editing = Submission.objects.filter(is_sent_for_editing=True).order_by('sent_for_edit_time')
    form = SubmissionForm()
    quarter = Quarter.objects.all()
    combined_content = "<br>".join([(submission.body) for submission in editing])

    if request.method == "POST":
        if 'remove' in request.POST:
            submission_id = request.POST.get('submission_id')
            if submission_id:
                submission = get_object_or_404(Submission, id=submission_id)
                submission.is_sent_for_editing = False
                submission.save()
                return redirect("editsubmissions")
            else:
                return redirect("editsubmissions")
        if 'reorder' in request.POST:
            new_order_submission_ids = request.POST.getlist('submission_id')
            current_time = timezone.now()
            for submission_id in new_order_submission_ids:
                submission = get_object_or_404(Submission, id=submission_id)
                submission.sent_for_edit_time = current_time
                submission.save()
            return redirect("editsubmissions")
        if 'publish_all' in request.POST:
            selected_submission_ids = request.POST.getlist('submission_id')
            form = NewsletterForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                for submission in editing:
                    if submission.is_sent_for_editing:
                        submission.is_sent_for_editing = False
                        submission.save()
                messages.success(request, 'New Newsletter Added')
                return redirect("newsletters")

    context = {
        'form': form,
        'editing':editing,
        'quarter':quarter,
        'combined_content': combined_content,
    }
    
    return render(request, 'editsubmissions.html', context)

@login_required(login_url='home')
def publishView(request):
    form = NewsletterForm()
    quarter = Quarter.objects.all()

    if request.method=="POST":
        form = NewsletterForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, 'New Newsletter Added')
            return redirect("newsletters")

    context = {'form':form,'quarter':quarter}
    return render(request,"publish.html", context)

def newslettersView(request):
    data = Newsletter.objects.all().order_by('upload_time')
    context = {'data':data}
    return render(request,'newsletters.html', context)

def newsletterView(request, pk):
    data = Newsletter.objects.get(id=pk)

    if request.method=="POST":
        data.delete()
        messages.error(request, 'Newsletter Deleted')
        return redirect("newsletters")

    context = {'data':data}
    return render(request,'newsletter.html', context)

@login_required(login_url='home')
def newsletterupdateView(request, pk):
    data = Newsletter.objects.get(id=pk)
    form = NewsletterForm(instance=data)
    quarter = Quarter.objects.all()
    is_editing = True
    
    if request.method=="POST":
        form = NewsletterForm(request.POST, request.FILES, instance=data)
        if form.is_valid():
            form.save()
            messages.success(request, 'Newsletter Updated')
            return redirect("/newsletter/"+pk)

    content = {'form':form, 'quarter':quarter, 'is_editing': is_editing}
    return render(request,"publish.html", content)

def searchView(request):
    form = SearchForm(request.GET)
    results = []
    if form.is_valid():
        search_query = form.cleaned_data['search_query']
        results = Newsletter.objects.filter(body__icontains=search_query)
    return render(request, 'newsletter.html', {'data': results})

def docxView(request, pk):
    newsletter = Newsletter.objects.get(id=pk)
    html_content = newsletter.body
    plain_text_content = html2text.html2text(html_content)
    doc = Document()
    doc.add_heading(f"Volume {newsletter.volume} | Issue {newsletter.issue}", 0)
    doc.add_paragraph(plain_text_content)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Newsletter_Volume_{newsletter.volume}_Issue_{newsletter.issue}.docx'
    doc.save(response)
    return response

@login_required(login_url='home')
def logoutView(request):
    logout(request)
    messages.error(request, 'Logged Out')
    return redirect('home')